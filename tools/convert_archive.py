#!/usr/bin/env python3
"""Convert freechecklists-archive non-PDF files to plain text.

Word-processing docs -> .txt, spreadsheets -> .tsv, HTML -> .txt. .zip archives
are extracted and their contents converted recursively. Images and PDFs are
*logged* (status "needs_ocr") for the OCR stage rather than converted here.

Output is a `text/` mirror tree of the archive's `checklists/` layout:
    checklists/<Mfr>/<Model>/<file>.doc  ->  text/<Mfr>/<Model>/<file>.txt

Dependencies (all optional at import time — missing ones log an error, don't crash):
    pip install xlrd python-docx openpyxl odfpy beautifulsoup4
    pandoc (for .rtf) + antiword binary for legacy .doc (Word 97-2003), with
    ANTIWORDHOME set to its data files (8859-*.txt etc.).

Usage:
    python3 tools/convert_archive.py
    python3 tools/convert_archive.py --archive /path/to/freechecklists-archive
    python3 tools/convert_archive.py --limit 20 --dry-run
"""
import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

# ── defaults ────────────────────────────────────────────────────────────────
ARCHIVE = Path("/home/node/workspace/openchecklists/freechecklists-archive")
CHECKLISTS = ARCHIVE / "checklists"
OUTPUT = ARCHIVE / "text"
MANIFEST = ARCHIVE / "convert-manifest.jsonl"
ANTIWORD = Path("/tmp/ocr-work/antiword-extract/usr/bin/antiword")
ANTIWORDHOME = Path("/tmp/ocr-work/antiword-extract/usr/share/antiword")

# Text-bearing source extensions (lowercase, no dot)
WORD_EXTS = {".doc", ".docx", ".wps", ".rtf", ".htm", ".html"}
SHEET_EXTS = {".xls", ".xlsx", ".ods", ".xlr"}
ZIP_EXTS = {".zip"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff"}
PLAIN_EXTS = {".txt", ".md", ".csv", ".diz", ".log", ".text"}
PDF_EXTS = {".pdf", ".pdf "}  # note trailing-space variant seen in the archive
SKIP_EXTS = {".odg", ".odp", ".pptx", ".ppt", ".dll", ".exe"}  # drawings/other


def classify(path: Path) -> str:
    """Classify by magic bytes where extension is ambiguous (.doc can be RTF)."""
    ext = path.suffix.lower()
    if ext in PDF_EXTS:
        return "pdf"
    if ext in ZIP_EXTS:
        return "zip"
    if ext in IMAGE_EXTS:
        return "image"
    if ext in PLAIN_EXTS:
        return "text"
    if ext in SKIP_EXTS:
        return "skip"
    if ext == ".doc":
        with open(path, "rb") as f:
            head = f.read(8)
        if head[:5].lower() == b"{\\rtf":
            return "rtf"
        return "doc_ole2"
    if ext in (".docx",):
        return "docx"
    if ext in (".rtf", ".wps"):
        return "rtf" if ext == ".rtf" else "wps"
    if ext in (".htm", ".html"):
        return "html"
    if ext in (".xls", ".xlsx", ".xlr"):
        return ext[1:]
    if ext in (".ods",):
        return "ods"
    return "unknown"


# ── converters (each returns str, or raises on failure) ─────────────────────

def _doc_ole2_to_text(path: Path) -> str:
    if not ANTIWORD.exists():
        raise RuntimeError(f"antiword not found at {ANTIWORD}")
    env = dict(os.environ)
    if ANTIWORDHOME.exists():
        env["ANTIWORDHOME"] = str(ANTIWORDHOME)
    out = subprocess.run(
        [str(ANTIWORD), str(path)],
        capture_output=True, text=True, encoding="utf-8", errors="replace", env=env, timeout=120,
    )
    if out.returncode != 0:
        raise RuntimeError(out.stderr.strip()[:300])
    return out.stdout


def _rtf_to_text(path: Path) -> str:
    out = subprocess.run(
        ["pandoc", "-f", "rtf", "-t", "plain", str(path)],
        capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=120,
    )
    if out.returncode != 0:
        raise RuntimeError(out.stderr.strip()[:300])
    return out.stdout


def _docx_to_text(path: Path) -> str:
    import docx  # noqa
    d = docx.Document(str(path))
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _html_to_text(path: Path) -> str:
    from bs4 import BeautifulSoup  # noqa
    soup = BeautifulSoup(path.read_bytes(), "html.parser")
    return soup.get_text("\n")


def _plain_to_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _cell_str(v) -> str:
    if v is None:
        return ""
    if isinstance(v, bool):
        return "TRUE" if v else "FALSE"
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


def _xls_to_tsv(path: Path) -> str:
    import xlrd  # noqa
    wb = xlrd.open_workbook(str(path), on_demand=True)
    out = []
    for sh in wb.sheets():
        out.append(f"# sheet: {sh.name}")
        for r in range(sh.nrows):
            cells = []
            for c in range(sh.ncols):
                cell = sh.cell(r, c)
                if cell.ctype == xlrd.XL_CELL_DATE:
                    try:
                        import datetime
                        cells.append(datetime.datetime(*xlrd.xldate_as_tuple(cell.value, wb.datemode)[:6]).strftime("%Y-%m-%d %H:%M"))
                        continue
                    except Exception:
                        pass
                cells.append(_cell_str(cell.value) if cell.ctype in (xlrd.XL_CELL_TEXT, xlrd.XL_CELL_NUMBER, xlrd.XL_CELL_BOOLEAN) else "")
            if any(cells):
                out.append("\t".join(cells))
    return "\n".join(out)


def _xlsx_to_tsv(path: Path) -> str:
    import openpyxl  # noqa
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [_cell_str(v) for v in row]
            if any(cells):
                out.append("\t".join(cells))
    return "\n".join(out)


def _ods_to_tsv(path: Path) -> str:
    import odf.opendocument  # noqa
    import odf.table  # noqa
    import odf.text  # noqa
    doc = odf.opendocument.load(str(path))
    out = []
    for table in doc.spreadsheet.getElementsByType(odf.table.Table):
        out.append(f"# sheet: {table.getAttribute('name')}")
        for row in table.getElementsByType(odf.table.TableRow):
            cells = []
            for cell in row.getElementsByType(odf.table.TableCell):
                text = "".join(
                    p.getAttribute("text") or ""
                    for p in cell.getElementsByType(odf.text.P)
                )
                # fall back to raw text nodes
                if not text:
                    text = "".join(
                        node.data for node in cell.childNodes if node.nodeType == node.TEXT_NODE
                    )
                cells.append(text.strip())
            if any(cells):
                out.append("\t".join(cells))
    return "\n".join(out)


CONVERTERS = {
    "doc_ole2": (_doc_ole2_to_text, ".txt"),
    "rtf": (_rtf_to_text, ".txt"),
    "docx": (_docx_to_text, ".txt"),
    "html": (_html_to_text, ".txt"),
    "text": (_plain_to_text, ".txt"),
    "xls": (_xls_to_tsv, ".tsv"),
    "xlsx": (_xlsx_to_tsv, ".tsv"),
    "ods": (_ods_to_tsv, ".tsv"),
}


def write_manifest(entry: dict, path: Path = MANIFEST):
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(entry) + "\n")


def convert_one(src: Path, out_dir: Path, manifest: Path, limit=None, counters=None) -> None:
    kind = classify(src)
    rel = src.relative_to(CHECKLISTS)
    entry = {"file": str(rel).replace("\\", "/"), "type": kind, "status": "pending"}

    if kind == "zip":
        _extract_zip(src, out_dir, manifest, counters)
        return
    if kind in ("pdf", "image"):
        entry["status"] = "needs_ocr"
        write_manifest(entry, manifest)
        return
    if kind in ("skip", "unknown", "wps", "xlr"):
        entry["status"] = "unsupported"
        write_manifest(entry, manifest)
        return

    conv, ext = CONVERTERS[kind]
    out_file = out_dir / rel.with_suffix(ext)
    if out_file.exists() and out_file.stat().st_size > 0:
        entry["status"] = "done"
        entry["output"] = str(out_file.relative_to(OUTPUT))
        write_manifest(entry, manifest)
        return

    try:
        text = conv(src)
    except Exception as e:
        entry["status"] = "error"
        entry["error"] = str(e)[:300]
        write_manifest(entry, manifest)
        return

    out_file.parent.mkdir(parents=True, exist_ok=True)
    out_file.write_text(text, encoding="utf-8")
    entry["status"] = "done"
    entry["output"] = str(out_file.relative_to(OUTPUT))
    write_manifest(entry, manifest)


def _extract_zip(zip_path: Path, out_dir: Path, manifest: Path, counters=None) -> None:
    rel = zip_path.relative_to(CHECKLISTS)
    base = out_dir / rel.with_suffix("")  # text/<Mfr>/<Model>/<zipname>/
    with tempfile.TemporaryDirectory() as td:
        tdir = Path(td)
        try:
            with zipfile.ZipFile(zip_path) as zf:
                zf.extractall(tdir)
        except Exception as e:
            write_manifest({"file": str(rel).replace("\\", "/"), "type": "zip", "status": "error", "error": str(e)[:300]}, manifest)
            return
        for inner in sorted(p for p in tdir.rglob("*") if p.is_file()):
            ikind = classify(inner)
            rel_inner = inner.relative_to(tdir)
            entry = {"file": f"{rel.as_posix()}::{rel_inner.as_posix()}", "type": ikind, "status": "pending"}
            if ikind in ("pdf", "image"):
                entry["status"] = "needs_ocr"
                write_manifest(entry, manifest)
                continue
            if ikind in CONVERTERS:
                conv, ext = CONVERTERS[ikind]
                out_file = base / rel_inner.with_suffix(ext)
                try:
                    text = conv(inner)
                    out_file.parent.mkdir(parents=True, exist_ok=True)
                    out_file.write_text(text, encoding="utf-8")
                    entry["status"] = "done"
                    entry["output"] = str(out_file.relative_to(OUTPUT))
                except Exception as e:
                    entry["status"] = "error"
                    entry["error"] = str(e)[:300]
                write_manifest(entry, manifest)
            else:
                entry["status"] = "unsupported"
                write_manifest(entry, manifest)


def main():
    global CHECKLISTS, OUTPUT, MANIFEST, ANTIWORD, ANTIWORDHOME
    ap = argparse.ArgumentParser(description="Convert non-PDF archive files to text")
    ap.add_argument("--archive", type=Path, default=ARCHIVE)
    ap.add_argument("--out", type=Path, default=OUTPUT)
    ap.add_argument("--manifest", type=Path, default=MANIFEST)
    ap.add_argument("--antiword", type=Path, default=ANTIWORD)
    ap.add_argument("--antiword-home", type=Path, default=ANTIWORDHOME)
    ap.add_argument("--limit", type=int, default=0, help="process at most N files")
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()

    CHECKLISTS = args.archive / "checklists"
    OUTPUT = args.out
    MANIFEST = args.manifest
    ANTIWORD = args.antiword
    ANTIWORDHOME = args.antiword_home

    files = sorted({p for p in CHECKLISTS.rglob("*") if p.is_file()})
    # Skip clean .pdf (already in the OCR batch). Keep ".pdf " (trailing-space)
    # and everything else so they're classified and logged (e.g. needs_ocr).
    files = [p for p in files if p.suffix.lower() != ".pdf"]
    if args.limit > 0:
        files = files[:args.limit]
    print(f"{len(files)} non-PDF files to process under {CHECKLISTS}")

    if args.dry_run:
        for p in files:
            print(f"  {classify(p):10} {p.relative_to(CHECKLISTS)}")
        return

    for i, src in enumerate(files, 1):
        convert_one(src, OUTPUT, MANIFEST)
        if i % 25 == 0:
            print(f"  ... {i}/{len(files)}")

    print("done")


if __name__ == "__main__":
    main()
